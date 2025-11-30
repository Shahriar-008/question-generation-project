import json
import os
import math
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION  # <-- NEW IMPORT

def to_bengali_numeral(num_str):
    """Converts an English number string to Bengali numerals."""
    english_to_bengali = {
        '0': '০', '1': '১', '2': '২', '3': '৩', '4': '৪',
        '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯'
    }
    bengali_str = ""
    for char in str(num_str):
        bengali_str += english_to_bengali.get(char, char)
    return bengali_str

# --- Helper function (No changes, but I fixed a small bug in your prefixes) ---
def add_question_to_document(document, q, show_answer=False):
    """Adds a single formatted question to the document object."""
    
    OMR_CIRCLE = '◯'
    FILLED_CIRCLE = '●'
    # I fixed this list, it previously contained incorrect characters
    bengali_prefixes = ['ক)', 'খ)', 'গ)', 'ঘ)'] 

    # A. Add Question Text
    p = document.add_paragraph()
    run = p.add_run(f'{q["id"]}. {q["question_text"]}')
    run.bold = True
    run.font.size = Pt(8)

    # B. Handle 'complex' questions
    if q["type"] == "complex":
        for sub_opt in q.get("sub_options", []):
            p_sub = document.add_paragraph(sub_opt, style='List Bullet')
            for run in p_sub.runs:
                run.font.size = Pt(8)
        prompt_p = document.add_paragraph()
        run_prompt = prompt_p.add_run(q.get("final_prompt", "নিচের কোনটি সঠিক?"))
        run_prompt.bold = True
        run_prompt.font.size = Pt(8)

    # C. Add Answer Options (in a 2x2 table)
    options = q.get("answer_options", [])
    correct_ans = str(q.get("correct_answer", "")).strip()

    if len(options) >= 4:
        opt_table = document.add_table(rows=2, cols=2)
        cell_data = [
            (opt_table.cell(0, 0), options[0], bengali_prefixes[0]),
            (opt_table.cell(0, 1), options[1], bengali_prefixes[1]),
            (opt_table.cell(1, 0), options[2], bengali_prefixes[2]),
            (opt_table.cell(1, 1), options[3], bengali_prefixes[3])
        ]
        for c, option_text, prefix in cell_data:
            p = c.paragraphs[0]
            p.text = '' 
            
            # Ensure strict boolean check and strip whitespace
            is_correct = (show_answer is True) and (str(option_text).strip() == correct_ans)
            current_circle = FILLED_CIRCLE if is_correct else OMR_CIRCLE

            run_circle = p.add_run(current_circle)
            run_circle.font.size = Pt(8)
            run_circle.font.name = 'Nirmala UI' 
            
            run_text = p.add_run(f' {prefix} {option_text}')
            run_text.font.size = Pt(8)
            run_text.font.name = 'Nirmala UI'
            
            if is_correct:
                run_text.bold = True
                run_circle.bold = True
    else:
        for j, opt in enumerate(options):
            is_correct = (show_answer is True) and (str(opt).strip() == correct_ans)
            p = document.add_paragraph()
            p.add_run(f'{bengali_prefixes[j]} {opt}')
            if is_correct:
                p.runs[0].bold = True
                p.add_run(' (সঠিক উত্তর)').bold = True

# --- Main function ---
def create_question_paper(json_file_path, output_docx_path, show_answer=False):
    """
    Generates a COMPACT two-column (snaking) MCQ paper
    with a FULL-WIDTH header.
    """
    
    # --- 1. Load Data ---
    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            questions = json.load(f)
    except FileNotFoundError:
        print(f"Error: JSON file not found at {json_file_path}")
        return
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON. Check for syntax errors in {json_file_path}")
        return

    # --- 2. Get Dynamic Counts ---
    num_questions = len(questions)
    bengali_marks = to_bengali_numeral(num_questions)
    bengali_time = bengali_marks 

    # --- 3. Initialize Document & Set Margins for Section 1 ---
    document = Document()
    section = document.sections[0] # This is the first, full-width section
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- 4. Set Global Font & Spacing Style ---
    style = document.styles['Normal']
    font = style.font
    font.name = 'Nirmala UI'
    font.size = Pt(10)
    
    p_format = style.paragraph_format
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(2)
    p_format.line_spacing = 1.0

    # --- 5. Add Header (to the first, full-width section) ---
    document.add_heading('Home Test', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('অষ্টম শ্রেণি (মাধ্যমিক) - ২০২৫').alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph('বিষয়: ').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_p = document.add_paragraph()
    p_format = header_p.paragraph_format
    tab_stops = p_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(
        section.page_width - section.left_margin - section.right_margin,
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    header_p.add_run(f'সময়— {bengali_time} মিনিট')
    header_p.add_run('\t')
    header_p.add_run(f'পূর্ণমান— {bengali_marks}')
    
    # --- 6. MODIFIED: Add NEW Section and set 2-Column Layout ---
    # Everything added AFTER this point will go into the new 2-column section
    
    new_section = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = new_section._sectPr # Get the new section's properties
    
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.insert_element_before(cols, 'w:docGrid')
    
    # Set to 2 columns and add 0.5 inch (720 twips) space
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '720')

    # --- 7. Populate Questions (into the new 2-column section) ---
    for q in questions:
        add_question_to_document(document, q, show_answer=show_answer)
        
    # --- 8. Save the Document ---
    try:
        output_dir = os.path.dirname(output_docx_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        document.save(output_docx_path)
        print(f"Successfully generated question paper with {num_questions} questions at: {output_docx_path}")
    
    except Exception as e:
        print(f"Error saving document: {e}")

# --- Main execution block ---
if __name__ == "__main__":
    
    input_file = "questions.json"
    output_folder = "output"
    
    # Clean up old file if it exists to avoid confusion
    old_file = os.path.join(output_folder, "Generated_Paper_Correct_Header.docx")
    if os.path.exists(old_file):
        try:
            os.remove(old_file)
            print(f"Removed old file: {old_file}")
        except OSError:
            pass

    # 1. Generate Student Copy (No Answers)
    print("Generating Student Copy (No Answers)...")
    output_file_student = os.path.join(output_folder, "Question_Paper.docx")
    create_question_paper(input_file, output_file_student, show_answer=False)

    # 2. Generate Teacher Copy (With Answers)
    print("Generating Teacher Copy (With Answers)...")
    output_file_teacher = os.path.join(output_folder, "Question_Paper_With_Answers.docx")
    create_question_paper(input_file, output_file_teacher, show_answer=True)